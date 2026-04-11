import os
from tools.registry import tool

@tool(
    name="read_pdf_content",
    description="Reads the text content of a PDF file.",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "The path to the PDF file."}
        },
        "required": ["file_path"]
    },
    requires_permission=True
)
def read_pdf(file_path: str) -> str:
    """Reads PDF and returns text."""
    try:
        from pypdf import PdfReader
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' not found."
            
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text if text.strip() else "Error: PDF seems empty or contains only non-extractable text (images?)."
    except ImportError:
        return "Error: 'pypdf' package is not installed. Please run: pip install pypdf"
    except Exception as e:
        return f"Error reading PDF: {e}"

@tool(
    name="read_docx_content",
    description="Reads the text content of a Word document (.docx).",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "The path to the .docx file."}
        },
        "required": ["file_path"]
    },
    requires_permission=True
)
def read_docx(file_path: str) -> str:
    """Reads DOCX and returns text."""
    try:
        from docx import Document
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' not found."
            
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        return "\n".join(full_text)
    except ImportError:
        return "Error: 'python-docx' package is not installed. Please run: pip install python-docx"
    except Exception as e:
        return f"Error reading Word document: {e}"

@tool(
    name="create_docx",
    description="Creates a new Word document (.docx) with the specified contents.",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "The path where the .docx file will be saved."},
            "content": {"type": "string", "description": "The text content to add to the document. Use double newlines for paragraphs."},
            "title": {"type": "string", "description": "Optional title heading for the document."}
        },
        "required": ["file_path", "content"]
    },
    requires_permission=True
)
def create_docx(file_path: str, content: str, title: str = None) -> str:
    """Creates a DOCX with the given content."""
    try:
        from docx import Document
        doc = Document()
        if title:
            doc.add_heading(title, 0)
        
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
                
        doc.save(file_path)
        return f"Successfully created document at '{file_path}'."
    except ImportError:
        return "Error: 'python-docx' package is not installed. Please run: pip install python-docx"
    except Exception as e:
        return f"Error creating Word document: {e}"
